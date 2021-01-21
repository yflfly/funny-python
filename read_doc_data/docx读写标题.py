# coding:utf-8
import docx

doc = docx.Document('data/1.docx')  # 打开文档

ps = doc.paragraphs
for p in ps:
    print(p.style)

re = []
for p in ps:
    if p.style.name == 'Heading 1':
        re.append((p.text, 1))

    if p.style.name == 'Heading 2':
        re.append((p.text, 2))
    if p.style.name == 'Heading 3':
        re.append((p.text, 3))
'''        
现在已经获取了标题内容以及标题的 level，将 re 列表“解压”：titles,titledes=zip(*re)，标题存在 titles 列表中，level 存在 titledes 列表中，接下来将标题写到新文档中
'''
print(re)
titles, titledes = zip(*re)
newdoc = docx.Document()
for i in range(len(titles)):
    newdoc.add_heading(titles[i], level=titledes[i])
newdoc.save('data/newfile.docx')
